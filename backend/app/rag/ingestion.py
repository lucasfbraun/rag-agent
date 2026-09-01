import os
import re
import uuid
import glob
import hashlib
from collections import defaultdict
from typing import Dict, List, Set, Tuple
from pypdf import PdfReader
import docx
from qdrant_client import QdrantClient
from qdrant_client.http import models as qmodels
from app.rag.embeddings import get_embedding
from app.config import QDRANT_HOST, QDRANT_PORT, COLLECTION_NAME, EMBEDDING_MODEL, VECTOR_SIZE

def get_qdrant_client() -> QdrantClient:
    """Cria e retorna um cliente Qdrant (lazy, evita falha no import-time)."""
    return QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, timeout=10)

def _garantir_indices_texto(client: QdrantClient) -> None:
    """Índices de texto tokenizado (palavra, minúsculas) nos campos `filename`
    e `content` — sem eles, o filtro MatchText usado na busca híbrida
    (app.rag.engine._detectar_codigos_produto / _extrair_palavras_chave) faz
    correspondência de substring bruta e case-sensitive (ex: buscar "2032"
    também bate em "203200"), que foi exatamente a causa de uma busca por
    código trazer produto errado. `content` cobre buscas por
    aplicação/uso quando a pergunta não cita um código de produto (ex: "cola
    para rolha de cortiça" não achava o FLEXX AG 2066 por embedding, mesmo o
    boletim tendo essas palavras quase literalmente). Idempotente: recriar um
    índice idêntico não é erro no Qdrant, seguro de chamar em toda ingestão."""
    for campo, min_token_len in (("filename", 2), ("content", 3)):
        try:
            client.create_payload_index(
                collection_name=COLLECTION_NAME,
                field_name=campo,
                field_schema=qmodels.TextIndexParams(
                    type="text",
                    tokenizer=qmodels.TokenizerType.WORD,
                    min_token_len=min_token_len,
                    lowercase=True,
                ),
            )
        except Exception as e:
            print(f"⚠️ Não foi possível garantir o índice de texto em '{campo}': {e}")


def init_qdrant_collection(client: QdrantClient):
    """Garante que a coleção de produtos e TDS exista no Qdrant."""
    collections = client.get_collections().collections
    if any(c.name == COLLECTION_NAME for c in collections):
        print(f"ℹ️ Coleção '{COLLECTION_NAME}' já existe — ingestão incremental.")
        _garantir_indices_texto(client)
        return
    try:
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=qmodels.VectorParams(
                size=VECTOR_SIZE,
                distance=qmodels.Distance.COSINE
            )
        )
        print(f"✅ Coleção '{COLLECTION_NAME}' inicializada.")
    except Exception as e:
        # Corrida entre duas ingestões simultâneas criando a coleção ao mesmo
        # tempo (AUD-007, achado novo). Se ela já existe agora, foi a outra
        # chamada vencendo a corrida — não é uma falha de verdade.
        colecoes_agora = client.get_collections().collections
        if any(c.name == COLLECTION_NAME for c in colecoes_agora):
            print(f"ℹ️ Coleção '{COLLECTION_NAME}' criada por outra ingestão concorrente.")
            _garantir_indices_texto(client)
            return
        print(f"❌ Erro ao inicializar coleção Qdrant: {e}")
        raise
    _garantir_indices_texto(client)

def extract_text_from_file(filepath: str) -> str:
    """Extrai texto e tabelas técnicas de PDFs, DOCX e TXT de produtos."""
    ext = os.path.splitext(filepath)[1].lower()
    text = ""

    if ext == ".pdf":
        reader = PdfReader(filepath)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"

    elif ext in (".docx", ".doc"):
        doc = docx.Document(filepath)
        for p in doc.paragraphs:
            text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                text += " | ".join(c.text.strip() for c in row.cells) + "\n"

    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

    return text

# Classificação de sensibilidade do RAG não estruturado (AUD-002, ticket 6).
# Deliberadamente ESTREITA — frases específicas, não palavras genéricas como
# "custo" sozinha — pra não classificar specs técnicas legítimas (densidade,
# NCO%, viscosidade) como sensíveis por engano; Vendedor precisa continuar
# vendo essas. Sub-inclusiva de propósito: baixo recall é um débito aceitável
# aqui, falso positivo bloqueando spec real não é. Ver docs/spec_rbac.md.
_PALAVRAS_CHAVE_SENSIVEIS = (
    "custo industrial",
    "custo unitário",
    "custo unitario",
    "preço de venda",
    "preco de venda",
    "margem de contribuição",
    "margem de contribuicao",
    "r$",
    "fórmula confidencial",
    "formula confidencial",
    "composição proprietária",
    "composicao proprietaria",
    "receita de formulação",
    "receita de formulacao",
    "matéria-prima e proporção",
    "materia-prima e proporcao",
)


def _e_conteudo_sensivel(texto: str) -> bool:
    """True se o chunk contém indício de custo/fórmula confidencial —
    controla `payload["sensivel"]` na ingestão e o filtro de
    `retrieve_products_context(..., incluir_sensivel=...)`."""
    texto_lower = texto.lower()
    return any(palavra in texto_lower for palavra in _PALAVRAS_CHAVE_SENSIVEIS)


# Deduplicação de arquivos (achado real: auditoria da coleção real mostrou
# ~4.900 chunks — quase metade do acervo — redundantes por dois padrões
# distintos, nenhum filtrado até esta sessão).

_EXTENSAO_PREFERENCIA = {".pdf": 0, ".docx": 1, ".doc": 2, ".txt": 3}

# "DOCUMENTAÇÃO DE PRODUTO RESTAURADO 0906" é uma árvore de pastas restaurada
# de um backup antigo que duplica, byte a byte, ~3.500 arquivos já existentes
# na árvore "atual" do acervo. Quando dois arquivos têm o MESMO CONTEÚDO
# (ver `_hash_conteudo`), o que estiver numa pasta com este marcador no
# caminho perde a disputa — mantém-se a cópia da árvore "viva".
_MARCADOR_PASTA_MENOS_PRIORITARIA = "restaurado"


def _agrupar_por_nome_base(arquivos: List[str]) -> Dict[Tuple[str, str], List[str]]:
    """Agrupa arquivos que ficam na MESMA pasta e têm o MESMO nome, só a
    extensão muda — achado real: ~1.200 pares de PDF/DOCX do mesmo boletim ou
    FISPQ indexados em dobro (a extração de texto difere um pouco entre
    formatos — ~95% de similaridade, não hash idêntico — por isso a detecção
    aqui é estrutural: mesma pasta, mesmo nome, não comparação de conteúdo)."""
    grupos: Dict[Tuple[str, str], List[str]] = defaultdict(list)
    for f in arquivos:
        pasta = os.path.dirname(f)
        base = os.path.splitext(os.path.basename(f))[0].strip().lower()
        grupos[(pasta, base)].append(f)
    return grupos


def _escolher_arquivo_preferido(caminhos: List[str]) -> str:
    """PDF > DOCX > DOC > TXT — PDF é o formato final/publicado mais comum
    no acervo; formatos não listados (não deveria acontecer, `supported` já
    filtra por extensão) ficam por último."""
    return min(caminhos, key=lambda c: _EXTENSAO_PREFERENCIA.get(os.path.splitext(c)[1].lower(), 99))


def _filtrar_duplicatas_de_formato(arquivos: List[str]) -> Tuple[List[str], List[str]]:
    """Retorna (mantidos, descartados): pra cada grupo de mesmo nome/pasta,
    mantém só 1 arquivo."""
    grupos = _agrupar_por_nome_base(arquivos)
    preferidos = {_escolher_arquivo_preferido(v) for v in grupos.values()}
    mantidos = [f for f in arquivos if f in preferidos]
    descartados = [f for f in arquivos if f not in preferidos]
    return mantidos, descartados


def _ordenar_por_prioridade(arquivos: List[str]) -> List[str]:
    """Processa primeiro os arquivos fora de pastas "restauradas" — em caso
    de empate de conteúdo idêntico (`_hash_conteudo`), o primeiro processado
    é o que fica indexado, os demais são ignorados como duplicata."""
    return sorted(arquivos, key=lambda c: (_MARCADOR_PASTA_MENOS_PRIORITARIA in c.lower(), c))


def _normalizar_texto(texto: str) -> str:
    return re.sub(r"\s+", " ", texto).strip().lower()


def _hash_conteudo(texto: str) -> str:
    """Hash do texto extraído, normalizado por espaço em branco — detecta
    arquivos com conteúdo EXATAMENTE idêntico (mesmo texto, caminhos
    diferentes), independente de nome de arquivo ou pasta."""
    return hashlib.sha256(_normalizar_texto(texto).encode("utf-8", errors="ignore")).hexdigest()


def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120) -> List[str]:
    """Divide texto em chunks com overlap para preservar contexto entre trechos."""
    words = text.split()
    chunks = []
    step = max(chunk_size - overlap, 1)
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i + chunk_size])
        if len(chunk.strip()) > 40:
            chunks.append(chunk)
    return chunks

def _pontos_existentes_por_arquivo(client: QdrantClient, collection_name: str) -> Dict[str, Set[str]]:
    """Varre a coleção inteira e devolve {filepath: {ponto_id, ...}}.

    Base da reconciliação (AUD-007, ticket 7): sem saber o que já está
    indexado por arquivo, não dá pra saber quais pontos ficaram obsoletos
    quando um arquivo encolhe, muda ou é removido do acervo."""
    existentes: Dict[str, Set[str]] = defaultdict(set)
    offset = None
    while True:
        pontos, offset = client.scroll(
            collection_name=collection_name,
            with_payload=["filepath"],
            with_vectors=False,
            limit=500,
            offset=offset,
        )
        for ponto in pontos:
            filepath = (ponto.payload or {}).get("filepath")
            if filepath:
                existentes[filepath].add(str(ponto.id))
        if offset is None:
            break
    return existentes


def _apagar_pontos(client: QdrantClient, collection_name: str, ids: Set[str]) -> None:
    if not ids:
        return
    client.delete(
        collection_name=collection_name,
        points_selector=qmodels.PointIdsList(points=list(ids)),
    )


def _arquivo_esta_no_escopo(filepath: str, dir_path_abs: str) -> bool:
    """True se `filepath` está dentro da árvore de `dir_path_abs`.

    Incidente real (Sessão 30): rodar a ingestão numa pasta PARCIAL (ex:
    `ingest_network.py --test`, que indexa só 1 família de produto, contra
    `--full`, o acervo inteiro) apagou a coleção inteira — tudo que não
    estava na pasta parcial foi tratado como "arquivo removido do acervo".
    A reconciliação de "arquivo removido" só pode considerar arquivos que
    ESTAVAM dentro do escopo desta execução; um filepath de fora nunca é
    candidato, não importa se apareceu ou não em `supported` desta vez."""
    try:
        file_abs = os.path.abspath(filepath)
        return os.path.commonpath([dir_path_abs, file_abs]) == dir_path_abs
    except ValueError:
        # Windows: caminhos em drives/raízes UNC diferentes levantam ValueError
        # no commonpath — nunca no mesmo escopo; na dúvida, não apaga.
        return False


def ingest_catalog_directory(dir_path: str, embedding_model: str = EMBEDDING_MODEL):
    """
    Indexa Boletins Técnicos (TDS), Catálogos e Homologações no Qdrant.

    - Idempotente: usa UUID como ID de ponto (não gera duplicatas em execuções repetidas
      para o mesmo conteúdo — IDs são hash determinístico do filepath+chunk_index).
    - Reconciliável (AUD-007, ticket 7): arquivo que encolheu, mudou ou saiu do
      acervo não deixa chunks obsoletos — a diferença entre o que já estava
      indexado por arquivo e o que a versão atual produz é apagada.
    - Suporte a PDF, DOCX e TXT.
    """
    client = get_qdrant_client()
    init_qdrant_collection(client)

    dir_path_abs = os.path.abspath(dir_path)
    existentes_por_arquivo = {
        filepath: ids
        for filepath, ids in _pontos_existentes_por_arquivo(client, COLLECTION_NAME).items()
        if _arquivo_esta_no_escopo(filepath, dir_path_abs)
    }

    files = glob.glob(os.path.join(dir_path, "**/*.*"), recursive=True)
    supported = [f for f in files if f.lower().endswith((".pdf", ".docx", ".doc", ".txt"))]

    # Deduplicação (achado real: auditoria da coleção mostrou ~4.900 chunks
    # redundantes — quase metade do acervo). Dois filtros ANTES de extrair
    # texto/gerar embedding, pra nem gastar esse custo com cópia descartada:
    # (1) mesmo nome+pasta, só a extensão muda — mantém 1 (_filtrar_duplicatas_de_formato);
    # (2) processa fora de pastas "restauradas" primeiro, pra elas perderem o
    # desempate de conteúdo idêntico (_ordenar_por_prioridade + hash abaixo).
    supported, descartados_formato = _filtrar_duplicatas_de_formato(supported)
    supported = _ordenar_por_prioridade(supported)

    print(f"🚀 Iniciando indexação de {len(supported)} arquivos de produtos (de {len(files)} total)...")
    if descartados_formato:
        print(
            f"⏭️ {len(descartados_formato)} arquivo(s) ignorado(s) por serem outro formato "
            "do mesmo documento já incluído (mesmo nome, mesma pasta)."
        )

    points = []
    indexed_chunks = 0
    skipped_files = 0
    removed_chunks = 0
    duplicados_de_conteudo = 0
    hashes_processados: Set[str] = set()

    for file_path in supported:
        filename = os.path.basename(file_path)
        # Pontos deste arquivo ficam numa lista própria até o arquivo inteiro
        # terminar com sucesso — uma falha no meio (embedding, extração) não
        # pode deixar chunks parciais gravados enquanto o arquivo é contado
        # como "não indexado" no relatório final (achado novo, ticket 7).
        pontos_do_arquivo = []
        ids_novos: Set[str] = set()
        try:
            raw_text = extract_text_from_file(file_path)
            if not raw_text.strip():
                print(f"⚠️ Arquivo vazio ou sem texto extraível: {filename}")
                skipped_files += 1
                continue

            hash_arquivo = _hash_conteudo(raw_text)
            if hash_arquivo in hashes_processados:
                # Não faz `existentes_por_arquivo.pop(file_path, ...)` de
                # propósito — se este arquivo já tinha pontos indexados de
                # uma execução anterior (antes deste filtro existir), eles
                # ficam "órfãos" e são removidos no passo final, igual um
                # arquivo removido do acervo.
                print(f"⏭️ '{filename}' ignorado: conteúdo idêntico a outro arquivo já indexado nesta execução.")
                duplicados_de_conteudo += 1
                continue
            hashes_processados.add(hash_arquivo)

            chunks = chunk_text(raw_text)

            for chunk_idx, chunk in enumerate(chunks):
                # ID determinístico: evita duplicatas se a ingestão for reexecutada
                point_uuid = str(uuid.uuid5(
                    uuid.NAMESPACE_URL,
                    f"{file_path}::{chunk_idx}"
                ))

                vector = get_embedding(chunk, embedding_model)

                pontos_do_arquivo.append(
                    qmodels.PointStruct(
                        id=point_uuid,
                        vector=vector,
                        payload={
                            "filename": filename,
                            "filepath": file_path,
                            "chunk_index": chunk_idx,
                            "content": chunk,
                            "sensivel": _e_conteudo_sensivel(chunk),
                        }
                    )
                )
                ids_novos.add(point_uuid)

        except Exception as e:
            print(f"⚠️ Erro ao processar '{filename}': {e}")
            skipped_files += 1
            continue

        points.extend(pontos_do_arquivo)
        indexed_chunks += len(pontos_do_arquivo)

        ids_antigos = existentes_por_arquivo.pop(file_path, set())
        obsoletos = ids_antigos - ids_novos
        if obsoletos:
            _apagar_pontos(client, COLLECTION_NAME, obsoletos)
            removed_chunks += len(obsoletos)

        if len(points) >= 100:
            client.upsert(collection_name=COLLECTION_NAME, points=points)
            points = []
            print(f"💾 {indexed_chunks} trechos indexados até agora...")

    if points:
        client.upsert(collection_name=COLLECTION_NAME, points=points)

    # Arquivos que estavam indexados mas não estão mais em `supported`
    # (removidos do disco, movidos, ou extensão deixou de ser suportada) —
    # os pontos restantes em existentes_por_arquivo são todos órfãos.
    for ids_orfaos in existentes_por_arquivo.values():
        _apagar_pontos(client, COLLECTION_NAME, ids_orfaos)
        removed_chunks += len(ids_orfaos)

    print(
        f"🎉 Indexação concluída! "
        f"{indexed_chunks} trechos de {len(supported) - skipped_files - duplicados_de_conteudo} arquivos indexados. "
        f"{skipped_files} arquivo(s) ignorado(s). "
        f"{duplicados_de_conteudo} duplicata(s) de conteúdo ignorada(s) nesta execução. "
        f"{len(descartados_formato)} duplicata(s) de formato ignorada(s) antes de processar. "
        f"{removed_chunks} chunk(s) obsoleto(s) removido(s)."
    )
